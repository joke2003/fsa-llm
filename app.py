import streamlit as st
import pandas as pd
import io
import json 
import os 
from datetime import datetime
from docx import Document 
import PyPDF2 
import concurrent.futures 
import re # Ensure re is imported

# --- Import from new modules ---
# These imports should NOT cause any Streamlit commands to run at import time.
from config import (
    APP_TITLE, APP_ICON, ANALYSIS_FRAMEWORK_SECTIONS, 
    ALL_DEFINED_MODULES_LIST, TOTAL_MODULES_COUNT,
    BASE_RESULT_DIR, PROMPTS_VERSION # Added PROMPTS_VERSION
)
# logger.py, llm_setup.py etc. are imported below after page_config

# --- !!! THIS MUST BE THE FIRST STREAMLIT COMMAND !!! ---
st.set_page_config(layout="wide", page_title=APP_TITLE, page_icon=APP_ICON)

# --- Now, import custom modules that might use Streamlit commands internally ---
# It's generally safer if these modules also avoid top-level st calls,
# but if they do (e.g. @st.cache_resource on a function), their import
# itself isn't the issue, it's when those functions are CALLED.

from logger import log_event
from llm_setup import get_llm_instance # Import the function, don't call it yet
from utils import (
    sanitize_filename, create_run_result_directory, get_latest_period_info,
    format_core_statements_for_llm, get_prior_analyses_summary 
)
from document_processing import preprocess_document_text 
from planning_services import (
    get_ai_planned_analysis_route, 
    plan_all_module_information_needs
)
from core_analysis_engine import run_llm_module_analysis 
from reporting import generate_and_save_html_report
from integration_services import consolidate_risks_and_opportunities, update_overall_conclusion_and_log_contradictions

# Try importing prompts.py after set_page_config, and handle error gracefully without st.error here
try:
    from prompts import MODULE_PROMPTS
except ImportError:
    # Log the error, but don't use st.error before the main app layout is built
    # This will be caught and displayed in the sidebar later if llm (and thus prompts) are needed.
    if 'run_log' not in st.session_state: st.session_state.run_log = [] # Ensure run_log exists
    log_event("CRITICAL_ERROR", "prompts.py 文件无法导入。应用功能将严重受限。", "AppSetup")
    MODULE_PROMPTS = {"DEFAULT_PROMPT": {"main_prompt_template": "错误：提示模块无法加载。"}}


# --- Core Working Paper (CWP) & Session State Initialization ---
def initialize_cwp():
    return {
        "base_data": { 
            "company_info": {
                "analysis_perspective": "股权投资", 
                "macro_analysis_conclusion_text": "用户未提供宏观经济分析结论，且未从测试目录加载默认文件。", 
                "industry_analysis_conclusion_text": "行业分析结论（基于波特五力模型）尚未生成。",
                "ai_planner_enabled": False 
            }, 
            "financial_reports": [] 
        },
        "analytical_module_outputs": {}, 
        "integrated_insights": { 
            "overall_summary": "", 
            "key_risks": [], 
            "key_opportunities": [],
            "current_overall_financial_conclusion": "分析尚未开始，暂无总体结论。", 
            "contradiction_logbook": [] 
        },
        "metadata_version_control": {
            "app_version": "0.10.2", # Corrected set_page_config execution order
            "analysis_timestamp": None, "llm_model_used": "DeepSeek (Conceptual)",
            "prompts_version": PROMPTS_VERSION, 
            "ai_planned_modules": None,
            "ai_planned_sections_for_display": None,
            "information_needs_by_module": {} 
        }
    }

# Initialize session state variables if they don't exist
if 'cwp' not in st.session_state: 
    st.session_state.cwp = initialize_cwp()
if 'run_log' not in st.session_state: # Must be initialized before get_llm_instance() if logger is used there
    st.session_state.run_log = [] 
if 'analysis_started' not in st.session_state: 
    st.session_state.analysis_started = False
if 'analysis_progress' not in st.session_state: 
    st.session_state.analysis_progress = 0
if 'current_module_processing' not in st.session_state: 
    st.session_state.current_module_processing = "等待开始..."
if 'num_periods_to_upload' not in st.session_state: 
    st.session_state.num_periods_to_upload = 1
if 'test_data_loaded_successfully' not in st.session_state: 
    st.session_state.test_data_loaded_successfully = False
if 'current_run_result_dir' not in st.session_state: 
    st.session_state.current_run_result_dir = None
if 'analysis_perspective_default' not in st.session_state: 
    st.session_state.analysis_perspective_default = "股权投资"
if 'ai_planner_toggle_default' not in st.session_state: 
    st.session_state.ai_planner_toggle_default = False

# --- Initialize LLM instance AFTER set_page_config and session_state init ---
llm = get_llm_instance() # This will trigger @st.cache_resource and st.secrets.get()


# --- Main UI and Workflow ---
st.title(f"{APP_ICON} {APP_TITLE}")
st.caption("基于“核心底稿”理念，利用大语言模型进行深度财务报表分析 (支持多期数据与趋势分析)。")
st.divider()

with st.sidebar:
    st.header("⚙️ 数据输入与设置")
    # Display LLM Initialization Status 
    if 'run_log' in st.session_state: # Check if run_log exists
        if llm is None: # Now check the initialized llm instance
            llm_init_error_messages = [entry['message'] for entry in st.session_state.run_log if entry.get('module') == "LLM_SETUP" and entry['type'] == "ERROR"]
            llm_init_warning_messages = [entry['message'] for entry in st.session_state.run_log if entry.get('module') == "LLM_SETUP" and entry['type'] == "WARNING"]
            if llm_init_error_messages:
                st.error(f"LLM 初始化失败: {llm_init_error_messages[0]}")
            elif llm_init_warning_messages:
                st.warning(f"LLM 配置问题: {llm_init_warning_messages[0]}")
            elif not any(entry.get('module') == "LLM_SETUP" for entry in st.session_state.run_log): # If no LLM_SETUP logs, means get_llm() might not have run or logged before failing
                 st.error("LLM 初始化状态未知或失败，请检查配置和日志。")
        else:
            st.success("LLM 已成功初始化。")
    else: # Should not happen if run_log is initialized above
        st.warning("日志系统未初始化。")


    with st.expander("1. 公司基本信息与分析角度", expanded=True):
        if 'company_name_default' not in st.session_state: st.session_state.company_name_default = "例如：贵州茅台股份有限公司"
        if 'industry_default' not in st.session_state: st.session_state.industry_default = "例如：白酒制造"
        if 'stock_code_default' not in st.session_state: st.session_state.stock_code_default = "例如：600519"
        if 'is_listed_default' not in st.session_state: st.session_state.is_listed_default = 0 
        
        company_name_input = st.text_input("公司名称*", value=st.session_state.company_name_default, key="company_name_input_key")
        is_listed_input = st.radio("是否上市公司*", ("是", "否"), index=st.session_state.is_listed_default, key="is_listed_radio_key")
        stock_code_input = st.text_input("股票代码 (如适用)", value=st.session_state.stock_code_default, key="stock_code_input_key")
        industry_input = st.text_input("所属行业*", value=st.session_state.industry_default, key="industry_input_key")
        analysis_perspective_options = ["股权投资", "债权投资", "债股双投"]
        analysis_perspective_input = st.selectbox("财务报表分析角度*", analysis_perspective_options, index=analysis_perspective_options.index(st.session_state.analysis_perspective_default), key="analysis_perspective_key")
        ai_planner_enabled_input = st.toggle("启用AI规划分析任务?", value=st.session_state.ai_planner_toggle_default, key="ai_planner_toggle_key",help="开启后，AI将根据公司信息和分析角度动态选择并排序分析模块。关闭则执行所有预设模块。")
    
    with st.expander("2. （可选）上传宏观经济分析结论", expanded=False):
        macro_analysis_file_input = st.file_uploader("上传宏观经济分析文件 (txt, md, pdf, docx)", type=['txt', 'md', 'pdf', 'docx'], key="macro_analysis_file_key")

    st.subheader("3. 上传财务报告期数据") 
    if st.session_state.get('test_data_loaded_successfully', False):
        st.success("一键测试数据已加载。您可修改上方公司信息后开始分析。")

    num_periods = st.number_input("选择上传报告期数量 (最多4期: 3年报+1季报)", min_value=1, max_value=4, value=st.session_state.num_periods_to_upload, key="num_periods_selector", disabled=st.session_state.get('test_data_loaded_successfully', False))
    if not st.session_state.get('test_data_loaded_successfully', False): st.session_state.num_periods_to_upload = num_periods

    uploaded_reports_data_sidebar = [] 
    for i in range(st.session_state.num_periods_to_upload):
        with st.container(): 
            st.markdown(f"##### 第 {i+1} 期报告数据")
            col_year, col_type = st.columns(2)
            with col_year: year_input_val = st.number_input(f"年份 (期 {i+1})*", min_value=2000, max_value=datetime.now().year + 1, value=datetime.now().year - i, key=f"year_{i}_key")
            with col_type: period_type_input_val = st.radio(f"报告类型 (期 {i+1})*", ("年报", "季报"), key=f"period_type_{i}_key", horizontal=True)
            quarter_input_val = None
            if period_type_input_val == "季报": quarter_input_val = st.selectbox(f"季度 (期 {i+1})*", (1, 2, 3, 4), format_func=lambda q: f"Q{q}", key=f"quarter_{i}_key")
            period_label_input_val = f"{year_input_val} {f'Q{quarter_input_val}' if quarter_input_val else 'Annual'}"
            st.caption(f"当前设定标签: {period_label_input_val}")
            bs_file_input_val = st.file_uploader(f"资产负债表 (期 {i+1})", type=['csv', 'xlsx'], key=f"bs_file_{i}_key")
            is_file_input_val = st.file_uploader(f"利润表 (期 {i+1})", type=['csv', 'xlsx'], key=f"is_file_{i}_key")
            cfs_file_input_val = st.file_uploader(f"现金流量表 (期 {i+1})", type=['csv', 'xlsx'], key=f"cfs_file_{i}_key")
            fn_file_input_val = st.file_uploader(f"财务报表附注 (期 {i+1}, 可选)", type=['txt', 'pdf', 'docx', 'md'], key=f"fn_file_{i}_key") 
            mda_file_input_val = st.file_uploader(f"管理层讨论与分析 (期 {i+1}, 可选)", type=['txt', 'pdf', 'docx', 'md'], key=f"mda_file_{i}_key") 
            uploaded_reports_data_sidebar.append({"year": year_input_val, "period_type": period_type_input_val, "quarter": quarter_input_val, "period_label": period_label_input_val, "bs_file": bs_file_input_val, "is_file": is_file_input_val, "cfs_file": cfs_file_input_val, "fn_file": fn_file_input_val, "mda_file": mda_file_input_val})
            if i < st.session_state.num_periods_to_upload -1 : st.markdown("---")

    col1_ctrl, col2_ctrl, col3_ctrl = st.columns(3) 
    with col1_ctrl: start_button = st.button("🚀 开始分析", type="primary", use_container_width=True, disabled=st.session_state.analysis_started)
    with col2_ctrl: reset_button = st.button("🔄 重置所有", use_container_width=True)
    with col3_ctrl: test_button = st.button("🧪 一键测试", use_container_width=True, help="从 ./test/ 目录加载预设的牧原股份报表文件及宏观分析文件。")

    if test_button:
        if not get_llm_instance(): 
            st.error("LLM未能成功初始化。请检查API密钥或相关配置。无法开始一键测试。")
            log_event("ERROR", "一键测试失败：LLM未初始化。", module_name="一键测试")
        else:
            st.session_state.cwp = initialize_cwp()
            st.session_state.run_log = []
            st.session_state.analysis_started = False 
            st.session_state.analysis_progress = 0
            st.session_state.current_module_processing = "一键测试数据加载中..."
            st.session_state.test_data_loaded_successfully = False 
            
            test_company_name = "牧原食品股份有限公司" 
            test_stock_code = "002714" 
            test_industry = "生猪养殖行业" 
            test_perspective = "股权投资" 
            test_ai_planner_enabled = False 
            
            st.session_state.current_run_result_dir = create_run_result_directory(test_company_name, BASE_RESULT_DIR)
            if not st.session_state.current_run_result_dir:
                 st.error("一键测试失败：无法创建结果目录。")
                 log_event("ERROR", "一键测试失败：无法创建结果目录。", module_name="一键测试")
            else:
                log_event("INFO", "一键测试数据加载开始。", module_name="一键测试")
                st.session_state.cwp['base_data']['company_info'] = {
                    "name": test_company_name, "is_listed": True,
                    "stock_code": test_stock_code, "industry": test_industry,
                    "analysis_date": pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S"),
                    "analysis_perspective": test_perspective, 
                    "ai_planner_enabled": test_ai_planner_enabled, 
                    "macro_analysis_conclusion_text": "一键测试：默认宏观经济分析结论。", 
                    "industry_analysis_conclusion_text": "" 
                }
                st.session_state.company_name_default = test_company_name
                st.session_state.industry_default = test_industry
                st.session_state.stock_code_default = test_stock_code
                st.session_state.is_listed_default = 0 
                st.session_state.analysis_perspective_default = test_perspective
                st.session_state.ai_planner_toggle_default = test_ai_planner_enabled
                log_event("CWP_INTERACTION", f"测试公司基本信息已写入核心底稿: {test_company_name}, 分析角度: {test_perspective}", module_name="一键测试")
                macro_filepath = os.path.join("test", "MACRO.md")
                if os.path.exists(macro_filepath):
                    try:
                        with open(macro_filepath, "r", encoding="utf-8") as f_macro: st.session_state.cwp['base_data']['company_info']['macro_analysis_conclusion_text'] = f_macro.read()
                        log_event("INFO", f"成功加载测试宏观分析文件: {macro_filepath}", module_name="一键测试")
                    except Exception as e: log_event("ERROR", f"加载测试宏观分析文件 {macro_filepath} 失败: {e}", module_name="一键测试"); st.session_state.cwp['base_data']['company_info']['macro_analysis_conclusion_text'] = "测试模式：加载 ./test/MACRO.md 文件时出错。"
                else: log_event("WARNING", f"测试宏观分析文件未找到: {macro_filepath}。将使用默认提示。", module_name="一键测试"); st.session_state.cwp['base_data']['company_info']['macro_analysis_conclusion_text'] = "测试模式：未在 ./test/ 目录找到 MACRO.md 文件。"
                
                test_data_path = "test"; years_to_test = [2023, 2022, 2021]; 
                file_map = {"BS": ("balance_sheet_data", "has_bs", ".xlsx"), "IS": ("income_statement_data", "has_is", ".xlsx"), "CFS": ("cash_flow_statement_data", "has_cfs", ".xlsx"), "NTS": ("footnotes_text_original", "has_fn", ".docx"), "MDA": ("mda_text_original", "has_mda", ".md")} 
                loaded_reports_count = 0; temp_reports_list = []
                
                for year in years_to_test:
                    period_label = f"{year} Annual (测试)"
                    period_entry = {"period_label": period_label, "year": year, "period_type": "年报", "quarter": None, "balance_sheet_data": None, "income_statement_data": None, "cash_flow_statement_data": None, "footnotes_text_original": "", "mda_text_original": "", "footnotes_processed_chunks": [], "mda_processed_chunks": [], "has_bs": False, "has_is": False, "has_cfs": False, "has_fn": False, "has_mda": False}
                    has_any_core_statement_for_year = False
                    for prefix, (data_key, has_key, ext) in file_map.items():
                        filename = f"{prefix}-{year}{ext}"; filepath = os.path.join(test_data_path, filename)
                        if os.path.exists(filepath):
                            try:
                                if ext == ".xlsx": 
                                    df = pd.read_excel(filepath); period_entry[data_key] = df.to_dict(); 
                                    if prefix in ["BS", "IS", "CFS"]: has_any_core_statement_for_year = True 
                                elif ext == ".docx":
                                    with open(filepath, "rb") as f_docx: doc = Document(io.BytesIO(f_docx.read())); full_text = [para.text for para in doc.paragraphs]; _=[full_text.append(f"\n--- 表格 {i+1} ---\n"+"\n".join(["\t|\t".join(c.text for c in r.cells) for r in t.rows])+"\n--- 表格结束 ---\n") for i,t in enumerate(doc.tables)]; period_entry[data_key] = "\n".join(full_text)
                                    if period_entry[data_key]: period_entry[f"{'footnotes' if prefix == 'NTS' else 'mda'}_processed_chunks"] = preprocess_document_text(period_entry[data_key], 'footnotes' if prefix == 'NTS' else 'mda', period_label)
                                elif ext == ".md" or ext == ".txt": 
                                    with open(filepath, "r", encoding="utf-8") as f_text: period_entry[data_key] = f_text.read()
                                    if period_entry[data_key]: period_entry[f"{'footnotes' if prefix == 'NTS' else 'mda'}_processed_chunks"] = preprocess_document_text(period_entry[data_key], 'footnotes' if prefix == 'NTS' else 'mda', period_label)
                                period_entry[has_key] = True; log_event("INFO", f"成功加载测试文件: {filepath}", module_name="一键测试")
                            except Exception as e: log_event("ERROR", f"加载或解析测试文件 {filepath} 失败: {e}", module_name="一键测试"); period_entry[has_key] = False 
                        else: log_event("WARNING", f"测试文件未找到: {filepath}", module_name="一键测试")
                    if has_any_core_statement_for_year: temp_reports_list.append(period_entry); loaded_reports_count += 1
                    else: log_event("WARNING", f"测试年份 {year} 无任何核心Excel报表文件，已跳过。", module_name="一键测试")
                
                st.session_state.cwp['base_data']['financial_reports'] = temp_reports_list
                st.session_state.cwp['base_data']['financial_reports'].sort(key=lambda x: (x['year'], x['quarter'] if x['period_type'] == '季报' else 0), reverse=True)
                st.session_state.cwp['metadata_version_control']['analysis_timestamp'] = pd.Timestamp.now().isoformat()
                st.session_state.cwp['metadata_version_control']['llm_model_used'] = "DeepSeek-Reasoner (Test Mode)"
                
                if loaded_reports_count > 0:
                    st.session_state.test_data_loaded_successfully = True; st.session_state.num_periods_to_upload = loaded_reports_count 
                    st.success(f"一键测试数据加载完毕 ({loaded_reports_count} 个报告期)。请在上方填写/确认公司信息与分析角度，然后点击“开始分析”。")
                    log_event("INFO", f"一键测试数据加载完成，共 {loaded_reports_count} 个报告期。", module_name="一键测试")
                else: st.error("一键测试未能加载任何报告期数据。请检查 `./test/` 目录下的文件。"); log_event("ERROR", "一键测试未能加载任何报告期数据。", module_name="一键测试")
                st.rerun()

    if start_button:
        current_company_name = company_name_input; current_is_listed = is_listed_input; current_stock_code = stock_code_input; current_industry = industry_input; current_analysis_perspective = analysis_perspective_input; current_macro_analysis_file = macro_analysis_file_input; current_ai_planner_enabled = ai_planner_enabled_input 
        if not current_company_name or not current_industry: error_msg = "请填写所有带 (*) 的必填项：公司名称和所属行业。"; st.error(error_msg); log_event("ERROR", f"开始分析失败：{error_msg}")
        elif not get_llm_instance(): error_msg = "LLM未能成功初始化。请检查API密钥或相关配置。无法开始分析。"; st.error(error_msg); log_event("ERROR", f"开始分析失败：{error_msg}")
        else:
            created_run_dir = create_run_result_directory(current_company_name, BASE_RESULT_DIR)
            if not created_run_dir: log_event("ERROR", "由于无法创建结果目录，分析流程中止。")
            else:
                st.session_state.current_run_result_dir = created_run_dir; st.session_state.analysis_started = True; st.session_state.analysis_progress = 0; st.session_state.current_module_processing = "数据预处理与规划中..." 
                ui_has_new_financial_report_files = any(prd.get("bs_file") or prd.get("is_file") or prd.get("cfs_file") or prd.get("fn_file") or prd.get("mda_file") for prd in uploaded_reports_data_sidebar)
                final_macro_text = st.session_state.cwp['base_data']['company_info'].get('macro_analysis_conclusion_text', "用户未提供宏观经济分析结论，且未从测试目录加载默认文件。")
                if current_macro_analysis_file:
                    log_event("INFO", f"开始处理用户上传的宏观经济分析文件: {current_macro_analysis_file.name}", module_name="数据预处理")
                    try:
                        if current_macro_analysis_file.name.endswith(".pdf"): pdf_reader = PyPDF2.PdfReader(current_macro_analysis_file); text = "".join(page.extract_text() for page in pdf_reader.pages if page.extract_text()); final_macro_text = text if text else f"PDF ({current_macro_analysis_file.name}) - No text extracted or empty."
                        elif current_macro_analysis_file.name.endswith(".docx"): doc = Document(current_macro_analysis_file);paras=[p.text for p in doc.paragraphs];tbls=[f"\nTable:\n"+"\n".join(["\t|\t".join(c.text for c in r.cells) for r in t.rows])+"\n--- Table End ---\n" for i,t in enumerate(doc.tables)];final_macro_text="\n".join(paras)+"\n".join(tbls) # type: ignore
                        elif current_macro_analysis_file.name.endswith((".txt",".md")): final_macro_text = current_macro_analysis_file.read().decode()
                        else: final_macro_text = f"Unsupported file type for Macro Analysis: {current_macro_analysis_file.name}"
                        if "Error reading" not in final_macro_text and "Unsupported file type" not in final_macro_text: log_event("INFO", f"用户上传的宏观经济分析文件 '{current_macro_analysis_file.name}' 处理完毕。", module_name="数据预处理")
                        else: log_event("ERROR", final_macro_text, module_name="数据预处理")
                    except Exception as e: final_macro_text = f"处理用户上传的宏观分析文件 '{current_macro_analysis_file.name}' 时发生意外错误: {e}"; log_event("ERROR", final_macro_text, module_name="数据预处理")
                
                if not st.session_state.test_data_loaded_successfully or ui_has_new_financial_report_files:
                    if ui_has_new_financial_report_files and st.session_state.test_data_loaded_successfully: log_event("INFO", "检测到用户通过UI上传了新文件，将优先处理UI文件，覆盖已加载的测试数据。", module_name="数据预处理")
                    elif ui_has_new_financial_report_files: log_event("INFO", "处理用户通过UI上传的文件。", module_name="数据预处理")
                    st.session_state.cwp = initialize_cwp(); st.session_state.run_log = [entry for entry in st.session_state.run_log if not (entry['type'] == 'CWP_INTERACTION' and ("测试公司" in entry.get('message', '') or "示例公司" in entry.get('message',''))) and not (entry.get('module_name') == '一键测试')]; st.session_state.test_data_loaded_successfully = False 
                    log_event("INFO", "分析流程开始 (用户触发 - 处理上传文件)。")
                    st.session_state.cwp['base_data']['company_info'] = {"name": current_company_name, "is_listed": current_is_listed == "是", "stock_code": current_stock_code if current_is_listed == "是" else "N/A", "industry": current_industry, "analysis_date": pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S"), "analysis_perspective": current_analysis_perspective, "ai_planner_enabled": current_ai_planner_enabled, "macro_analysis_conclusion_text": final_macro_text, "industry_analysis_conclusion_text": "行业分析结论（基于波特五力模型）尚未生成。"}
                    processed_reports_from_ui = []
                    try:
                        for report_data in uploaded_reports_data_sidebar:
                            if not (report_data["bs_file"] and report_data["is_file"] and report_data["cfs_file"]):
                                if any([report_data["bs_file"], report_data["is_file"], report_data["cfs_file"], report_data["fn_file"], report_data["mda_file"]]): log_event("WARNING", f"报告期 {report_data['period_label']} 缺少核心三表，将跳过此期。", module_name="数据预处理")
                                continue
                            period_entry = {"period_label": report_data["period_label"], "year": report_data["year"], "period_type": report_data["period_type"], "quarter": report_data["quarter"], "balance_sheet_data": None, "income_statement_data": None, "cash_flow_statement_data": None, "footnotes_text_original": "", "mda_text_original": "", "footnotes_processed_chunks": [], "mda_processed_chunks": [], "has_bs": False, "has_is": False, "has_cfs": False, "has_fn": False, "has_mda": False}
                            if report_data["bs_file"]: period_entry["balance_sheet_data"] = (pd.read_excel(report_data["bs_file"]) if report_data["bs_file"].name.endswith('xlsx') else pd.read_csv(report_data["bs_file"])).to_dict(); period_entry["has_bs"] = True
                            if report_data["is_file"]: period_entry["income_statement_data"] = (pd.read_excel(report_data["is_file"]) if report_data["is_file"].name.endswith('xlsx') else pd.read_csv(report_data["is_file"])).to_dict(); period_entry["has_is"] = True
                            if report_data["cfs_file"]: period_entry["cash_flow_statement_data"] = (pd.read_excel(report_data["cfs_file"]) if report_data["cfs_file"].name.endswith('xlsx') else pd.read_csv(report_data["cfs_file"])).to_dict(); period_entry["has_cfs"] = True
                            for file_type_key, cwp_text_key_orig, cwp_chunks_key, doc_type_name, has_key in [
                                ("fn_file", "footnotes_text_original", "footnotes_processed_chunks", "footnotes", "has_fn"), 
                                ("mda_file", "mda_text_original", "mda_processed_chunks", "mda", "has_mda")
                            ]:
                                uploaded_file = report_data[file_type_key]
                                if uploaded_file:
                                    file_content = ""
                                    if uploaded_file.name.endswith(".pdf"):
                                        try: pdf_reader = PyPDF2.PdfReader(uploaded_file); text = "".join(page.extract_text() for page in pdf_reader.pages if page.extract_text()); file_content = text if text else f"PDF ({uploaded_file.name}) - No text extracted or empty."
                                        except Exception as e: file_content = f"Error reading PDF {uploaded_file.name}: {e}"; log_event("ERROR", file_content, module_name="数据预处理")
                                    elif uploaded_file.name.endswith(".docx"):
                                        try: doc = Document(uploaded_file); full_text = [para.text for para in doc.paragraphs];_=[full_text.append(f"\n--- 表格 {i+1} ---\n"+"\n".join(["\t|\t".join(c.text for c in r.cells) for r in t.rows])+"\n--- 表格结束 ---\n") for i,t in enumerate(doc.tables)];file_content="\n".join(full_text) # type: ignore
                                        except Exception as e: file_content = f"Error reading DOCX {uploaded_file.name}: {e}"; log_event("ERROR", file_content, module_name="数据预处理")
                                    elif uploaded_file.name.endswith((".txt", ".md")):
                                        try: file_content = uploaded_file.read().decode()
                                        except Exception as e: file_content = f"Error reading {uploaded_file.name.split('.')[-1].upper()}: {e}"; log_event("ERROR", file_content, module_name="数据预处理")
                                    else: file_content = f"Unsupported file type: {uploaded_file.name}"; log_event("WARNING", file_content, module_name="数据预处理")
                                    
                                    period_entry[cwp_text_key_orig] = file_content 
                                    if file_content and not file_content.startswith("Error"):
                                        st.session_state.current_module_processing = f"预处理文档: {uploaded_file.name} ({period_entry['period_label']})..."
                                        period_entry[cwp_chunks_key] = preprocess_document_text(file_content, doc_type_name, period_entry['period_label'])
                                        period_entry[has_key] = True
                                    else: period_entry[has_key] = False
                            processed_reports_from_ui.append(period_entry); log_event("CWP_INTERACTION", f"UI上传的报告期 {report_data['period_label']} 数据已处理并预分块。", module_name="数据预处理")
                        
                        st.session_state.cwp['base_data']['financial_reports'] = processed_reports_from_ui
                        st.session_state.cwp['base_data']['financial_reports'].sort(key=lambda x: (x['year'], x['quarter'] if x['period_type'] == '季报' else 0), reverse=True)
                        if not st.session_state.cwp['base_data']['financial_reports']: st.error("未能成功处理任何通过UI上传的报告期数据。请确保至少一个报告期包含核心三表。"); log_event("ERROR", "未能成功处理任何通过UI上传的报告期数据。", module_name="数据预处理"); st.session_state.analysis_started = False 
                        else: log_event("INFO", f"成功处理 {len(st.session_state.cwp['base_data']['financial_reports'])} 期来自UI的报告数据。", module_name="数据预处理")
                    except Exception as e: st.error(f"处理UI上传文件失败: {e}"); log_event("ERROR", f"处理UI上传文件失败: {e}", module_name="数据预处理"); st.session_state.analysis_started = False 
                else: 
                    log_event("INFO", "分析流程开始 (用户触发 - 使用已加载的测试数据)。")
                    st.session_state.cwp['base_data']['company_info'].update({ "name": current_company_name, "is_listed": current_is_listed == "是", "stock_code": current_stock_code if current_is_listed == "是" else "N/A", "industry": current_industry, "analysis_date": pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S"), "analysis_perspective": current_analysis_perspective, "ai_planner_enabled": current_ai_planner_enabled, "macro_analysis_conclusion_text": final_macro_text })

                log_event("CWP_INTERACTION", f"公司基本信息已写入/更新核心底稿: {st.session_state.cwp['base_data']['company_info']['name']}, 分析角度: {st.session_state.cwp['base_data']['company_info']['analysis_perspective']}, AI规划器: {'启用' if st.session_state.cwp['base_data']['company_info']['ai_planner_enabled'] else '关闭'}", module_name="数据预处理")
                if st.session_state.cwp['base_data']['company_info']['macro_analysis_conclusion_text'] and "用户未提供" not in st.session_state.cwp['base_data']['company_info']['macro_analysis_conclusion_text'] and "失败" not in st.session_state.cwp['base_data']['company_info']['macro_analysis_conclusion_text']:
                    log_event("CWP_INTERACTION", "宏观经济分析结论已存入核心底稿。", module_name="数据预处理", details={"snippet": st.session_state.cwp['base_data']['company_info']['macro_analysis_conclusion_text'][:100]+"..."})
                else: log_event("WARNING", f"宏观经济分析结论最终为: {st.session_state.cwp['base_data']['company_info']['macro_analysis_conclusion_text']}", module_name="数据预处理")
                
                if st.session_state.analysis_started: 
                    modules_for_info_planning = []
                    if st.session_state.cwp['base_data']['company_info'].get('ai_planner_enabled', False):
                        if st.session_state.cwp['metadata_version_control'].get('ai_planned_modules') is None: 
                            st.session_state.current_module_processing = "AI任务规划器运行中..."
                            planned_route = get_ai_planned_analysis_route(
                                st.session_state.cwp['base_data']['company_info'],
                                st.session_state.cwp['base_data']['company_info'].get('macro_analysis_conclusion_text', ''),
                                ALL_DEFINED_MODULES_LIST
                            )
                            st.session_state.cwp['metadata_version_control']['ai_planned_modules'] = planned_route
                        modules_for_info_planning = st.session_state.cwp['metadata_version_control'].get('ai_planned_modules', []) 
                        if not modules_for_info_planning: 
                            modules_for_info_planning = ALL_DEFINED_MODULES_LIST
                            log_event("WARNING", "AI规划模块列表为空，将为所有预定义模块规划信息需求。", "InfoNeedsPlannerTrigger")
                    else:
                        modules_for_info_planning = ALL_DEFINED_MODULES_LIST
                    
                    if modules_for_info_planning:
                        st.session_state.current_module_processing = "批量规划模块信息需求中..."
                        log_event("INFO", "开始批量规划所有待执行模块的信息需求。", "InfoNeedsPlannerTrigger")
                        
                        available_docs_summary_for_planner = "当前已加载的、可供提取详细信息的文档包括：\n"
                        for r_idx, r_entry in enumerate(st.session_state.cwp['base_data']['financial_reports']):
                            available_docs_summary_for_planner += f"- 报告期: {r_entry['period_label']}: "
                            doc_types = []
                            if r_entry.get("footnotes_processed_chunks"): doc_types.append(f"财务报表附注 (共 {len(r_entry['footnotes_processed_chunks'])} 块)")
                            if r_entry.get("mda_processed_chunks"): doc_types.append(f"管理层讨论与分析 (共 {len(r_entry['mda_processed_chunks'])} 块)")
                            available_docs_summary_for_planner += ", ".join(doc_types) if doc_types else "无已处理的补充文档分块"
                            available_docs_summary_for_planner += "\n"

                        information_needs_plan = plan_all_module_information_needs(
                            modules_to_plan_for=modules_for_info_planning,
                            company_info=st.session_state.cwp['base_data']['company_info'],
                            macro_conclusion=st.session_state.cwp['base_data']['company_info'].get('macro_analysis_conclusion_text', ''),
                            industry_conclusion=st.session_state.cwp['base_data']['company_info'].get('industry_analysis_conclusion_text', ''),
                            available_docs_summary=available_docs_summary_for_planner
                        )
                        st.session_state.cwp['metadata_version_control']['information_needs_by_module'] = information_needs_plan
                        log_event("CWP_INTERACTION", "所有模块的信息需求规划结果已存入核心底稿。", "InfoNeedsPlannerTrigger", {"num_modules_planned": len(information_needs_plan)})
                    else:
                         log_event("WARNING", "没有模块需要进行信息需求规划 (模块列表为空)。", "InfoNeedsPlannerTrigger")

                    st.session_state.cwp['metadata_version_control']['analysis_timestamp'] = pd.Timestamp.now().isoformat()
                    st.session_state.cwp['metadata_version_control']['llm_model_used'] = "DeepSeek-Reasoner"
                    st.rerun()

    if reset_button:
        # ... (Unchanged)
        st.session_state.cwp = initialize_cwp()
        st.session_state.analysis_started = False; st.session_state.analysis_progress = 0
        st.session_state.current_module_processing = ""; st.session_state.num_periods_to_upload = 1
        st.session_state.run_log = []; st.session_state.test_data_loaded_successfully = False
        st.session_state.company_name_default = "例如：贵州茅台股份有限公司"; st.session_state.industry_default = "例如：白酒制造"
        st.session_state.stock_code_default = "例如：600519"; st.session_state.is_listed_default = 0
        st.session_state.analysis_perspective_default = "股权投资"; st.session_state.ai_planner_toggle_default = False
        st.session_state.current_run_result_dir = None 
        st.success("所有输入和分析结果已重置。"); st.rerun()

    if st.session_state.analysis_started:
        st.sidebar.progress(st.session_state.analysis_progress / 100, text=st.session_state.current_module_processing)

# --- Main UI Logic ---
if not st.session_state.analysis_started:
    st.info("请在左侧侧边栏输入公司信息并上传财务报表（支持多期），然后点击“开始分析”。或点击“一键测试”从本地 `./test/` 目录加载预设文件。")
else:
    modules_to_execute_ordered = []
    sections_to_execute_for_display = {} 
    TOTAL_MODULES_TO_RUN_CURRENT = 0 
    
    if st.session_state.cwp['base_data']['company_info'].get('ai_planner_enabled', False):
        planned_modules = st.session_state.cwp['metadata_version_control'].get('ai_planned_modules') 
        if planned_modules: 
            modules_to_execute_ordered = planned_modules
            temp_sections = {}
            for sec, mods_in_framework in ANALYSIS_FRAMEWORK_SECTIONS.items():
                current_sec_mods = [m for m in mods_in_framework if m in planned_modules]
                if current_sec_mods:
                    temp_sections[sec] = sorted(current_sec_mods, key=lambda x: planned_modules.index(x) if x in planned_modules else float('inf'))
            sections_to_execute_for_display = temp_sections
            TOTAL_MODULES_TO_RUN_CURRENT = len(modules_to_execute_ordered)
            st.session_state.cwp['metadata_version_control']['ai_planned_sections_for_display'] = sections_to_execute_for_display 
        else: 
            log_event("WARNING", "AI规划器启用但未返回有效模块计划，将执行所有预定义模块。", "MainLogic")
            modules_to_execute_ordered = ALL_DEFINED_MODULES_LIST 
            sections_to_execute_for_display = ANALYSIS_FRAMEWORK_SECTIONS
            TOTAL_MODULES_TO_RUN_CURRENT = TOTAL_MODULES_COUNT 
            st.session_state.cwp['metadata_version_control']['ai_planned_sections_for_display'] = ANALYSIS_FRAMEWORK_SECTIONS
    else: 
        modules_to_execute_ordered = ALL_DEFINED_MODULES_LIST 
        sections_to_execute_for_display = ANALYSIS_FRAMEWORK_SECTIONS
        TOTAL_MODULES_TO_RUN_CURRENT = TOTAL_MODULES_COUNT 
        st.session_state.cwp['metadata_version_control']['ai_planned_sections_for_display'] = ANALYSIS_FRAMEWORK_SECTIONS 
    
    if not st.session_state.cwp['metadata_version_control'].get('information_needs_by_module') and modules_to_execute_ordered:
        st.session_state.current_module_processing = "首次运行时批量规划模块信息需求中..."
        log_event("INFO", "首次运行或AI规划模块更新，开始批量规划信息需求。", "MainLogic")
        available_docs_summary_for_planner = "可查询文档包括：\n"
        for r_idx, r_entry in enumerate(st.session_state.cwp['base_data']['financial_reports']):
            available_docs_summary_for_planner += f"- 报告期: {r_entry['period_label']}: "
            doc_types = []
            if r_entry.get("footnotes_processed_chunks"): doc_types.append(f"财务报表附注 (共 {len(r_entry['footnotes_processed_chunks'])} 块)")
            if r_entry.get("mda_processed_chunks"): doc_types.append(f"管理层讨论与分析 (共 {len(r_entry['mda_processed_chunks'])} 块)")
            available_docs_summary_for_planner += ", ".join(doc_types) if doc_types else "无已处理的补充文档分块"
            available_docs_summary_for_planner += "\n"
        information_needs_plan = plan_all_module_information_needs(
            modules_to_plan_for=modules_to_execute_ordered,
            company_info=st.session_state.cwp['base_data']['company_info'],
            macro_conclusion=st.session_state.cwp['base_data']['company_info'].get('macro_analysis_conclusion_text', ''),
            industry_conclusion=st.session_state.cwp['base_data']['company_info'].get('industry_analysis_conclusion_text', ''),
            available_docs_summary=available_docs_summary_for_planner
        )
        st.session_state.cwp['metadata_version_control']['information_needs_by_module'] = information_needs_plan
        log_event("CWP_INTERACTION", "所有模块的信息需求规划结果已存入核心底稿。", "MainLogic", {"num_modules_planned": len(information_needs_plan)})
        st.rerun() 

    analysis_in_progress_main = len(st.session_state.cwp['analytical_module_outputs']) < TOTAL_MODULES_TO_RUN_CURRENT

    if analysis_in_progress_main and st.session_state.cwp['metadata_version_control'].get('information_needs_by_module'): 
        st.header("🚀 分析进行中...")
        col_prog_bar, col_prog_text = st.columns([3, 1])
        with col_prog_bar: st.progress(st.session_state.analysis_progress / 100 if st.session_state.analysis_progress is not None and TOTAL_MODULES_TO_RUN_CURRENT > 0 else 0)
        with col_prog_text: st.write(f"{st.session_state.analysis_progress if st.session_state.analysis_progress is not None else 0}% 完成")
        st.info(f"当前正在处理模块: **{st.session_state.current_module_processing}**")
        completed_module_names = list(st.session_state.cwp['analytical_module_outputs'].keys())
        pending_module_names = [m for m in modules_to_execute_ordered if m not in completed_module_names]
        st.markdown("---"); col_completed, col_pending = st.columns(2)
        with col_completed:
            st.subheader(f"已完成模块 ({len(completed_module_names)}/{TOTAL_MODULES_TO_RUN_CURRENT}):")
            if completed_module_names: 
                for mod_item_val_c in completed_module_names: 
                    if isinstance(mod_item_val_c, str) and mod_item_val_c:
                        st.markdown(f"- ✅ {mod_item_val_c}") 
                    else: 
                        log_event("ERROR", f"Skipping display: Invalid item in completed_module_names. Type: {type(mod_item_val_c)}, Value: '{mod_item_val_c}'", "UI_Error_Display")
            else: st.caption("暂无已完成模块。")
        with col_pending:
            st.subheader("待处理模块:")
            if pending_module_names: 
                for mod_item_val_p in pending_module_names: 
                    if isinstance(mod_item_val_p, str) and mod_item_val_p:
                        st.markdown(f"- ⏳ {mod_item_val_p}") 
                    else:
                        log_event("ERROR", f"Skipping display: Invalid item in pending_module_names. Type: {type(mod_item_val_p)}, Value: '{mod_item_val_p}'", "UI_Error_Display")
            else: st.caption("所有规划模块已加入处理队列或已完成。")
        st.markdown("---")

        next_module_to_run = None; next_section_key_for_module = None
        for mod_name in modules_to_execute_ordered:
            if mod_name not in st.session_state.cwp['analytical_module_outputs']:
                next_module_to_run = mod_name
                for sec, mods_in_sec in ANALYSIS_FRAMEWORK_SECTIONS.items(): 
                    if mod_name in mods_in_sec: next_section_key_for_module = sec; break
                break
        if next_module_to_run and next_section_key_for_module:
            run_llm_module_analysis(next_module_to_run, next_section_key_for_module)
            st.rerun() 
        elif not pending_module_names and len(completed_module_names) >= TOTAL_MODULES_TO_RUN_CURRENT: 
             st.session_state.current_module_processing = "✅ 分析全部完成！"
             log_event("INFO", "所有规划的分析模块已完成（主循环检测）。")
             if not st.session_state.cwp['integrated_insights'].get('key_risks') and TOTAL_MODULES_TO_RUN_CURRENT > 0 : consolidate_risks_and_opportunities() 
             if not st.session_state.cwp['integrated_insights'].get('overall_summary') and TOTAL_MODULES_TO_RUN_CURRENT > 0: generate_and_save_html_report()
             st.rerun()
    elif not analysis_in_progress_main: # Analysis is complete
        tab_titles = ["⚙️ 运行日志", "📊 总览与摘要", "🌍 战略与环境", "📈 业绩与效率", "💰 盈利与会计", "📉 风险与偿债", "🚀 增长与持续", "🔮 预测与建模", "⚖️ 公司估值", "📝 核心底稿追踪"]
        tabs = st.tabs(tab_titles)
        with tabs[0]: 
            st.header("⚙️ 运行日志"); st.markdown("按时间倒序显示系统运行过程中的关键事件。")
            if st.button("刷新日志", key="refresh_log_button_main_final"): st.rerun()
            if not st.session_state.run_log: st.info("暂无运行日志。请开始分析以生成日志。")
            else:
                log_container = st.container(height=600); 
                for entry in st.session_state.run_log: 
                    log_class = f"log-entry log-entry-{entry['type']}"; module_prefix = f"模块: `{entry['module']}` - " if 'module' in entry else "系统 - "; message_display = entry['message']
                    if entry.get('details'):
                        details_str_parts = []; 
                        if 'query' in entry['details']: details_str_parts.append(f"查询: `{entry['details']['query']}`")
                        if 'result_snippet' in entry['details']: details_str_parts.append(f"结果摘要: `{entry['details']['result_snippet']}`")
                        if 'summary_snippet' in entry['details']: details_str_parts.append(f"摘要: `{entry['details']['summary_snippet']}`")
                        if 'inventory' in entry['details']: details_str_parts.append(f"文档清单: {entry['details']['inventory']}")
                        if 'prompt_length' in entry['details']: details_str_parts.append(f"提示长度: {entry['details']['prompt_length']}")
                        if 'tool_name' in entry['details']: details_str_parts.append(f"工具: `{entry['details']['tool_name']}`")
                        if 'args' in entry['details']: args_display = entry['details']['args'];_ = json.dumps(args_display, ensure_ascii=False) if isinstance(args_display, dict) else str(args_display); details_str_parts.append(f"参数: `{_}`")
                        if 'conversation' in entry['details']: 
                             try: 
                                 conv_str = json.dumps(entry['details']['conversation'], ensure_ascii=False, indent=2)
                                 details_str_parts.append(f"对话历史: \n```json\n{conv_str}\n```")
                             except TypeError: details_str_parts.append("对话历史: (序列化失败)")
                        elif 'full_prompt' in entry['details']: 
                             details_str_parts.append(f"完整提示: \n```\n{entry['details']['full_prompt']}\n```")

                        if details_str_parts: message_display += " (" + ", ".join(details_str_parts) + ")"
                    with log_container: st.markdown(f"<div class='{log_class}'><b>{entry['timestamp']} [{entry['type']}]</b> {module_prefix}{message_display}</div>", unsafe_allow_html=True)
        with tabs[1]:
            st.header("📊 分析总览与核心结论摘要")
            if st.session_state.analysis_progress >= 100 and llm: 
                if st.session_state.cwp['integrated_insights'].get('key_risks'):
                    st.subheader("主要风险点")
                    for risk_idx, risk in enumerate(st.session_state.cwp['integrated_insights']['key_risks']):
                        with st.expander(f"风险: {risk.get('description', 'N/A')[:50]}... (ID: {risk.get('id', 'N/A')})", expanded=False): 
                            st.markdown(f"**描述:** {risk.get('description', 'N/A')}"); st.markdown(f"**分类:** {risk.get('category', 'N/A')}"); st.markdown(f"**潜在影响:** {risk.get('potential_impact', 'N/A')}")
                            if risk.get('source_modules'): st.markdown(f"**来源模块:** {', '.join(risk['source_modules'])}")
                            if risk.get('mitigating_factors_observed'): st.markdown(f"**缓解因素:** {risk['mitigating_factors_observed']}")
                            if risk.get('notes_for_further_investigation'): st.markdown(f"**进一步调查:** {risk['notes_for_further_investigation']}")
                if st.session_state.cwp['integrated_insights'].get('key_opportunities'):
                    st.subheader("主要机遇点")
                    for opp_idx, opp in enumerate(st.session_state.cwp['integrated_insights']['key_opportunities']):
                         with st.expander(f"机遇: {opp.get('description', 'N/A')[:50]}... (ID: {opp.get('id', 'N/A')})", expanded=False): 
                            st.markdown(f"**描述:** {opp.get('description', 'N/A')}"); st.markdown(f"**分类:** {opp.get('category', 'N/A')}"); st.markdown(f"**潜在收益:** {opp.get('potential_benefit', 'N/A')}")
                            if opp.get('source_modules'): st.markdown(f"**来源模块:** {', '.join(opp['source_modules'])}")
                            if opp.get('actionability_notes'): st.markdown(f"**行动建议/关注点:** {opp['actionability_notes']}")
                st.divider(); st.subheader("最终总体财务分析摘要")
                cwp_module_summaries = []; ordered_module_names_final = [mod_name for section_modules in ANALYSIS_FRAMEWORK_SECTIONS.values() for mod_name in section_modules] 
                for mod_name_sum in ordered_module_names_final:
                    if mod_name_sum in st.session_state.cwp['analytical_module_outputs']:
                        mod_output_sum = st.session_state.cwp['analytical_module_outputs'][mod_name_sum]
                        if mod_output_sum.get('status') == 'Completed':
                            summary_to_use = mod_output_sum.get('abbreviated_summary') or mod_output_sum.get('text_summary', ''); confidence_to_use = mod_output_sum.get('confidence_score', 'N/A')
                            cwp_module_summaries.append(f"模块 '{mod_name_sum}' (置信度: {confidence_to_use}): {summary_to_use[:200]}...")
                report_labels_summary = ', '.join([r['period_label'] for r in st.session_state.cwp['base_data']['financial_reports']])
                _, local_latest_period_label_summary = get_latest_period_info(st.session_state.cwp) 
                final_current_overall_conclusion = st.session_state.cwp['integrated_insights'].get('current_overall_financial_conclusion', "未能生成迭代的总体结论。")
                contradiction_log_content = "无记录的矛盾点。"; 
                if st.session_state.cwp['integrated_insights'].get('contradiction_logbook'):
                    contradiction_log_content = "\n分析过程中记录的潜在矛盾点：\n"; 
                    for item_idx, item in enumerate(st.session_state.cwp['integrated_insights']['contradiction_logbook']): contradiction_log_content += f"{item_idx+1}. 模块'{item['module_name']}' (置信度: {item['module_confidence']}) 指出: {item['contradiction_description']}\n"
                macro_conclusion_for_summary = st.session_state.cwp['base_data']['company_info'].get('macro_analysis_conclusion_text', '用户未提供宏观经济分析结论。')
                industry_conclusion_for_summary = st.session_state.cwp['base_data']['company_info'].get('industry_analysis_conclusion_text', '行业分析结论尚未生成。')
                summary_prompt_template = f"""
                您是一位资深的财务分析师...撰写一份全面且富有洞察力的最终总体财务分析摘要。
                **公司名称:** {st.session_state.cwp['base_data']['company_info']['name']} **所属行业:** {st.session_state.cwp['base_data']['company_info']['industry']} **分析角度:** {st.session_state.cwp['base_data']['company_info'].get('analysis_perspective', '未指定')}
                **已分析的报告期包括:** {report_labels_summary}. **最新报告期为:** {local_latest_period_label_summary}.
                **A. 用户提供的宏观经济分析结论：**\n```{macro_conclusion_for_summary}```
                **B. 系统生成的行业分析结论：**\n```{industry_conclusion_for_summary}```
                **C. 最终的“（当前）公司总体财务分析结论”：**\n```{final_current_overall_conclusion}```
                **D. 分析过程中记录的“矛盾点记录本”：**\n```{contradiction_log_content}```
                **E. 已识别的关键风险点：**\n```{json.dumps(st.session_state.cwp['integrated_insights'].get('key_risks', []), ensure_ascii=False, indent=2)}```
                **F. 已识别的关键机遇点：**\n```{json.dumps(st.session_state.cwp['integrated_insights'].get('key_opportunities', []), ensure_ascii=False, indent=2)}```
                **G. 各个独立分析模块的摘要：**\n```{ "\n".join(cwp_module_summaries)}```
                **撰写要求：** 1. 全面性与深度（2000-3000汉字）。2. 整合性（体现分析角度，结合宏观行业）。3. 关注矛盾点、风险与机遇。4. 结构建议（宏观行业背景、公司概况、财务核心、优势机遇、风险挑战、增长持续、估值、总结展望）。
                请生成最终总体财务分析摘要。
                """
                try:
                    summary_messages = [{"role": "user", "content": summary_prompt_template}]; summary_response = llm.invoke(summary_messages) 
                    final_summary = summary_response.content if hasattr(summary_response, 'content') else str(summary_response)
                    st.session_state.cwp['integrated_insights']['overall_summary'] = final_summary; st.markdown(final_summary)
                    log_event("CWP_INTERACTION", "最终总体财务分析摘要已生成并存入核心底稿。", module_name="总览与摘要")
                except Exception as e: st.error(f"生成最终摘要失败: {e}"); log_event("ERROR", f"生成最终摘要失败: {e}", module_name="总览与摘要")
            elif not llm and st.session_state.analysis_started: st.warning("LLM 未初始化，无法生成最终摘要。")
            elif st.session_state.analysis_started: st.info("所有分析模块完成后，将在此处生成总体摘要。")
            else: st.info("请先开始分析。")

        # Display other analysis module tabs
        ordered_sections_for_display = sections_to_execute_for_display 
        tab_offset = 2 
        current_tab_idx = 0
        for section_title_key, modules_in_section in ordered_sections_for_display.items():
            if not modules_in_section: continue 
            original_section_index = -1
            for idx, (orig_sec_key, _) in enumerate(ANALYSIS_FRAMEWORK_SECTIONS.items()):
                if orig_sec_key == section_title_key: original_section_index = idx; break
            if original_section_index != -1:
                display_tab_index = original_section_index + tab_offset
                if display_tab_index < len(tabs) -1: 
                    with tabs[display_tab_index]:
                        st.header(f"第 {original_section_index + 1} 部分：{section_title_key}") 
                        st.markdown(f"*本部分包含 {len(modules_in_section)} 个分析模块。*")
                        for module_idx, module_name in enumerate(modules_in_section): 
                            with st.expander(f"**{module_name}**", expanded=False): 
                                output_data = st.session_state.cwp['analytical_module_outputs'].get(module_name)
                                if output_data:
                                    st.markdown(f"**置信度:** {output_data.get('confidence_score', 'N/A')}") 
                                    st.markdown(output_data.get("text_summary", "无文本摘要。")) 
                                    if output_data.get("structured_data"): st.caption("结构化数据:"); st.json(output_data["structured_data"], expanded=False)
                                    st.caption(f"状态: {output_data.get('status', 'N/A')} | 时间: {output_data.get('timestamp', 'N/A')}")
                                    col_prompt, col_messages = st.columns(2)
                                    with col_prompt:
                                        if st.toggle("显示使用的提示", key=f"toggle_prompt_{module_name}_{original_section_index}_{display_tab_index}_{module_idx}", value=False): 
                                            st.text_area("Prompt Used", value=output_data.get("prompt_used", "无提示信息记录。"), height=200, disabled=True, key=f"prompt_disp_{module_name}_{original_section_index}_{display_tab_index}_{module_idx}")
                                    with col_messages:
                                         if output_data.get("message_history") and st.toggle("显示交互历史", key=f"toggle_messages_{module_name}_{original_section_index}_{display_tab_index}_{module_idx}", value=False): 
                                            st.json(output_data.get("message_history"), expanded=False)
                                else: st.info(f"模块 '{module_name}' 分析结果待生成或未被AI规划器选中。")
            current_tab_idx +=1
        
        with tabs[-1]: 
            st.header("📝 核心底稿实时追踪")
            with st.expander("1. 基础数据层", expanded=False):
                st.subheader("公司基本信息"); st.json(st.session_state.cwp['base_data']['company_info'], expanded=True)
                st.subheader("已上传财务报告期数据")
                if st.session_state.cwp['base_data']['financial_reports']:
                    for i, report_entry in enumerate(st.session_state.cwp['base_data']['financial_reports']):
                        with st.container(border=True):
                            st.markdown(f"**报告期 {i+1}: {report_entry['period_label']} ({report_entry['period_type']})**")
                            st.caption(f"年份: {report_entry['year']}" + (f", 季度: Q{report_entry['quarter']}" if report_entry['quarter'] else ""))
                            docs_present = []; 
                            if report_entry.get("has_bs"): docs_present.append("资产负债表")
                            if report_entry.get("has_is"): docs_present.append("利润表")
                            if report_entry.get("has_cfs"): docs_present.append("现金流量表")
                            if report_entry.get("footnotes_processed_chunks"): docs_present.append(f"附注 ({len(report_entry['footnotes_processed_chunks'])} 块)") 
                            if report_entry.get("mda_processed_chunks"): docs_present.append(f"MD&A ({len(report_entry['mda_processed_chunks'])} 块)") 
                            st.write(f"已上传/处理文件: {', '.join(docs_present) if docs_present else '无核心文件或未处理'}")
                            
                            if report_entry.get("footnotes_processed_chunks"):
                                with st.popover(f"查看 {report_entry['period_label']} 附注分块概述", use_container_width=True):
                                    for chunk_idx, chunk_data in enumerate(report_entry["footnotes_processed_chunks"]):
                                        st.markdown(f"**块 {chunk_idx+1} (ID: {chunk_data['chunk_id']}) 概述:**")
                                        st.caption(chunk_data['overview_text'])
                                        st.divider()
                            if report_entry.get("mda_processed_chunks"):
                                with st.popover(f"查看 {report_entry['period_label']} MD&A分块概述", use_container_width=True):
                                    for chunk_idx, chunk_data in enumerate(report_entry["mda_processed_chunks"]):
                                        st.markdown(f"**块 {chunk_idx+1} (ID: {chunk_data['chunk_id']}) 概述:**")
                                        st.caption(chunk_data['overview_text'])
                                        st.divider()
                else: st.caption("尚未上传或处理任何报告期数据。")
            
            with st.expander("迭代总结与矛盾点", expanded=True):
                st.subheader("当前公司总体财务分析结论 (迭代更新)")
                st.markdown(st.session_state.cwp['integrated_insights'].get('current_overall_financial_conclusion', '暂无'))
                st.subheader("矛盾点记录本")
                if st.session_state.cwp['integrated_insights'].get('contradiction_logbook'):
                    for idx, entry in enumerate(st.session_state.cwp['integrated_insights']['contradiction_logbook']):
                        st.markdown(f"**矛盾点 {idx+1}:**")
                        st.markdown(f"- **记录时间:** {entry['timestamp']}")
                        st.markdown(f"- **引发模块:** {entry['module_name']} (置信度: {entry['module_confidence']})")
                        st.markdown(f"- **矛盾描述:** {entry['contradiction_description']}")
                        with st.popover(f"查看矛盾点 {idx+1} 相关结论片段", key=f"popover_contradiction_{idx}_cwp_key"):
                            st.markdown(f"**模块结论片段:**\n```\n{entry['module_finding_snippet']}\n```")
                            st.markdown(f"**前期总体结论片段:**\n```\n{entry['previous_overall_conclusion_snippet']}\n```")
                        st.divider()
                else:
                    st.markdown("暂无记录的矛盾点。")
            
            with st.expander("关键风险与机遇", expanded=True):
                st.subheader("识别出的关键风险点")
                if st.session_state.cwp['integrated_insights'].get('key_risks'):
                    for risk_idx_cwp, risk in enumerate(st.session_state.cwp['integrated_insights']['key_risks']):
                        st.markdown(f"- **{risk.get('description')}** (分类: {risk.get('category','N/A')}, 影响: {risk.get('potential_impact','N/A')})", key=f"cwp_risk_{risk_idx_cwp}")
                else:
                    st.markdown("暂未汇总关键风险。")
                
                st.subheader("识别出的关键机遇点")
                if st.session_state.cwp['integrated_insights'].get('key_opportunities'):
                    for opp_idx_cwp, opp in enumerate(st.session_state.cwp['integrated_insights']['key_opportunities']):
                        st.markdown(f"- **{opp.get('description')}** (分类: {opp.get('category','N/A')}, 收益: {opp.get('potential_benefit','N/A')})", key=f"cwp_opp_{opp_idx_cwp}")
                else:
                    st.markdown("暂未汇总关键机遇。")


            with st.expander("分析模块输出层", expanded=False):
                if st.session_state.cwp['analytical_module_outputs']:
                    for module_name, output_val in st.session_state.cwp['analytical_module_outputs'].items():
                        with st.container(border=True):
                            st.markdown(f"##### {module_name}")
                            st.markdown(f"**状态:** {output_val.get('status', 'N/A')} | **时间:** {output_val.get('timestamp', 'N/A')} | **置信度:** {output_val.get('confidence_score', 'N/A')}")
                            if output_val.get("text_summary"): st.caption("文本摘要:"); st.markdown(f"> {output_val['text_summary'][:300]}...")
                            if output_val.get("abbreviated_summary"): st.caption("缩略摘要:"); st.markdown(f"> {output_val['abbreviated_summary'][:300]}...") 
                            if output_val.get("structured_data"): st.caption("结构化数据:"); st.json(output_val["structured_data"], expanded=False)
                            col_prompt_cwp, col_messages_cwp = st.columns(2)
                            with col_prompt_cwp:
                                if output_val.get("prompt_used") and st.toggle("显示使用的提示", key=f"toggle_prompt_cwp_{module_name}_cwp_disp_key_v5", value=False): 
                                     st.text_area("Prompt Used (CWP)", value=output_val["prompt_used"], height=150, disabled=True, key=f"prompt_cwp_disp_{module_name}_cwp_area_key_v5")
                            with col_messages_cwp:
                                if output_val.get("message_history") and st.toggle("显示交互历史", key=f"toggle_messages_cwp_{module_name}_cwp_disp_key_v5", value=False): 
                                    st.json(output_val.get("message_history"), expanded=False)
                else: st.info("分析模块的输出将在此处逐条记录。")
            with st.expander("最终综合洞察与结论层 (摘要)", expanded=False): 
                st.markdown(f"**最终总体分析摘要:**"); st.markdown(f"{st.session_state.cwp['integrated_insights'].get('overall_summary', '分析完成后生成。')}")
            with st.expander("元数据与版本控制层", expanded=False): 
                st.json(st.session_state.cwp['metadata_version_control'])

